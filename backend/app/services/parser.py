from __future__ import annotations

import io
import logging
from pathlib import Path

import docx
from pypdf import PdfReader
from pypdf.errors import PyPdfError

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx"}
MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024  # 10 MB


class ParserError(Exception):
    """Raised when a file cannot be parsed into text."""


def _detect_extension(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ParserError(
            f"Unsupported file format: {ext or 'unknown'}. "
            f"Allowed: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )
    return ext


def _parse_txt(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp1251", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ParserError("Failed to decode txt file with utf-8/cp1251/latin-1")


def _parse_pdf(data: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(data))
    except (PyPdfError, OSError, ValueError) as exc:
        raise ParserError(f"Failed to open PDF: {exc}") from exc

    if reader.is_encrypted:
        raise ParserError("Encrypted PDFs are not supported")

    parts: list[str] = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception as exc:
            logger.warning("pdf page extraction failed: %s", exc)
            text = ""
        if text:
            parts.append(text)

    result = "\n".join(parts).strip()
    if not result:
        raise ParserError("PDF contains no extractable text (likely scanned image)")
    return result


def _parse_docx(data: bytes) -> str:
    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise ParserError(f"Failed to open DOCX: {exc}") from exc

    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    result = "\n".join(parts).strip()
    if not result:
        raise ParserError("DOCX contains no extractable text")
    return result


def parse(filename: str, data: bytes) -> tuple[str, str]:
    """Parse uploaded file bytes into text.

    Returns (text, source_format) where source_format is one of "txt"/"pdf"/"docx".
    Raises ParserError on any failure with a human-readable message.
    """
    if not data:
        raise ParserError("Empty file")
    if len(data) > MAX_FILE_SIZE_BYTES:
        raise ParserError(
            f"File too large: {len(data)} bytes (max {MAX_FILE_SIZE_BYTES} bytes)"
        )

    ext = _detect_extension(filename)
    fmt = ext.lstrip(".")

    if ext == ".txt":
        return _parse_txt(data), fmt
    if ext == ".pdf":
        return _parse_pdf(data), fmt
    if ext == ".docx":
        return _parse_docx(data), fmt

    raise ParserError(f"Unsupported format: {ext}")
