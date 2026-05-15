"""Smoke test for parser: txt/pdf/docx + edge cases. Read-only verification."""
from __future__ import annotations

import io
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "backend"))

from app.services.parser import ParserError, parse  # noqa: E402


def main() -> int:
    # txt UTF-8
    text, fmt = parse("test.txt", "Привет мир\nЭто тест парсера".encode("utf-8"))
    assert "Привет" in text and fmt == "txt", text
    print(f"OK txt utf-8         len={len(text):3d} fmt={fmt}")

    # txt cp1251 fallback
    text, fmt = parse("cp.txt", "Привет".encode("cp1251"))
    assert "Привет" in text and fmt == "txt"
    print(f"OK txt cp1251        len={len(text):3d} fmt={fmt}")

    # docx
    import docx as docx_lib

    doc = docx_lib.Document()
    doc.add_paragraph("Резюме: Иван Иванов")
    doc.add_paragraph("Python, Django, PostgreSQL")
    buf = io.BytesIO()
    doc.save(buf)
    text, fmt = parse("resume.docx", buf.getvalue())
    assert "Иван" in text and "Python" in text and fmt == "docx"
    print(f"OK docx              len={len(text):3d} fmt={fmt}")

    # unsupported extension
    try:
        parse("x.exe", b"binary")
    except ParserError as exc:
        print(f"OK unsupported       -> {exc}")
    else:
        return 1

    # bad pdf bytes
    try:
        parse("garbage.pdf", b"not really a pdf")
    except ParserError as exc:
        print(f"OK bad pdf           -> {exc}")
    else:
        return 1

    # empty file
    try:
        parse("empty.txt", b"")
    except ParserError as exc:
        print(f"OK empty             -> {exc}")
    else:
        return 1

    # too large
    try:
        parse("huge.txt", b"x" * (11 * 1024 * 1024))
    except ParserError as exc:
        print(f"OK oversize          -> {exc}")
    else:
        return 1

    print("\nALL parser checks passed")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
